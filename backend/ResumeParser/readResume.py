import os
import fitz  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

def readResume(file_path):
    """
    Extract text from various resume file formats
    Supports: PDF, DOCX, DOC, TXT
    """
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None
    
    # Get file extension
    file_extension = Path(file_path).suffix.lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        elif file_extension == '.doc':
            return extract_text_from_doc(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file format: {file_extension}")
            return None
            
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return None

def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            
        doc.close()
        return text.strip()
        
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None

def extract_text_from_doc(file_path):
    """
    Extract text from DOC files using python-docx2txt
    Note: This is a fallback method for older .doc files
    """
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        return text.strip() if text else None
        
    except ImportError:
        print("docx2txt not installed. Install with: pip install docx2txt")
        return extract_text_from_doc_alternative(file_path)
    except Exception as e:
        print(f"Error extracting text from DOC: {e}")
        return None

def extract_text_from_doc_alternative(file_path):
    """
    Alternative method for DOC files using antiword (if available)
    """
    try:
        import subprocess
        result = subprocess.run(['antiword', file_path], 
                              capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout.strip()
        else:
            print("Could not extract text from DOC file. Consider converting to DOCX.")
            return None
    except FileNotFoundError:
        print("antiword not found. For better DOC support, install antiword or convert to DOCX.")
        return None
    except Exception as e:
        print(f"Error with antiword: {e}")
        return None

def extract_text_from_txt(file_path):
    """Extract text from TXT files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content.strip()
        
    except UnicodeDecodeError:
        # Try different encodings
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                return content.strip()
            except UnicodeDecodeError:
                continue
        
        print(f"Could not decode text file with any encoding")
        return None
        
    except Exception as e:
        print(f"Error reading text file: {e}")
        return None

def get_supported_formats():
    return ['.pdf', '.docx', '.doc', '.txt']

def is_supported_format(file_path):
    file_extension = Path(file_path).suffix.lower()
    return file_extension in get_supported_formats()